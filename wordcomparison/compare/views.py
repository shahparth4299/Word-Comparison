from django.http import HttpResponse
from django.shortcuts import render
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from django.utils.encoding import smart_text
import docx
import diff_match_patch
#index view
def index(request):
	return render(request,'index.html')

#new view
def new(request):
	return HttpResponse("Welcome to the New Page")

#upload_file view 
def upload_file(request):

	file1_string = handle_file(request.FILES['file1'])
	file2_string = handle_file(request.FILES['file2'])
	
	diff_obj = WordComparison()
	result = diff_obj.diff_main(file1_string, file2_string)
	diff_obj.diff_cleanupSemantic(result)

	file1_changes = diff_obj.file1_changes(result)
	file2_changes = diff_obj.file2_changes(result)

	context = {"file1_changes":file1_changes,"file2_changes":file2_changes}
	
	return render(request, 'display.html',context)

class WordComparison(diff_match_patch.diff_match_patch):
	def file1_changes(self, diffs):
		html = []
		for (flag, data) in diffs:
			text = (data.replace("&", "&amp;")
				.replace("<", "&lt;")
				.replace(">", "&gt;")
				.replace("\n", "<br>"))
			if flag == self.DIFF_DELETE:
				html.append("""<b style=\"background:#ffe6e6;
					\">%s</b>""" % text)
			elif flag == self.DIFF_EQUAL:
				html.append("<span>%s</span>" % text)
		return "".join(html)

	def file2_changes(self, diffs):
		html = []
		for (flag, data) in diffs:
			text = (data.replace("&", "&amp;")
				.replace("<", "&lt;")
				.replace(">", "&gt;")
				.replace("\n", "<br>"))
			if flag == self.DIFF_INSERT:
				html.append("""<b style=\"background:#e6ffe6;
					\">%s</b>""" % text)
			elif flag == self.DIFF_EQUAL:
				html.append("<span>%s</span>" % text)
		return "".join(html)

def handle_file(f1):
	output = ""
	if(".docx" in f1.name):
		doc = docx.Document(f1)
		for x in doc.paragraphs:
			token1 = sent_tokenize(x.text)
			for j in token1:
				output += j
	elif(".txt" in f1.name):
		for x in f1:
			temp = x.decode('utf-8')
			temp2 = temp.replace("\n", "").replace("\r", "")
			token1 = sent_tokenize(temp2)
			for j in token1:
				output += j
	return output

#Manual Comparison
def manual_comparison(request):
	
	read_file1 = handle_uploaded_file(request.FILES['file1'])
	read_file2 = handle_uploaded_file(request.FILES['file2'])		

	file1_parser = get_list(request.FILES['file1'])
	file2_parser = get_list(request.FILES['file2'])

	dissimilar_lines1 = perform_comparision(file1_parser,file2_parser)
	dissimilar_words = find_words(request.FILES['file1'],request.FILES['file2'])

	dissimilar_lines2 = perform_comparision(file2_parser,file1_parser)

	output_text1 = get_edited_data(request.FILES['file1'],dissimilar_lines1)
	output_text2 = get_edited_data(request.FILES['file2'],dissimilar_lines2)
	
	context = {'file_content1': read_file1,'file_content2': read_file2,"dissimilar_lines1":dissimilar_lines1,"dissimilar_words":dissimilar_words,"output_text1":output_text1,"output_text2":output_text2}
	
	return render(request, 'display.html',context)


def get_list(f1):
	file_docs = []
	if(".docx" in f1.name):
		doc = docx.Document(f1)
		for i in doc.paragraphs:
			token1 = sent_tokenize(i.text)
			for j in token1 :
				file_docs.append(j)
	elif(".txt" in f1.name):
		for x in f1:
			temp = x.decode('utf-8')
			temp2 = temp.replace("\n", "").replace("\r", "")
			token1 = sent_tokenize(temp2)
			for j in token1:
				file_docs.append(j)
	return file_docs


def handle_uploaded_file(f1):
    output = ""
    if(".docx" in f1.name):
    	doc = docx.Document(f1)
    	for x in doc.paragraphs:
    		token1 = sent_tokenize(x.text)
    		for j in token1:
    			output += j
    elif(".txt" in f1.name):
    	for x in f1:
    		temp = x.decode('utf-8')
    		temp2 = temp.replace("\n", "").replace("\r", "")
    		token1 = sent_tokenize(temp2)
    		for j in token1:
    			output + j
    return output

def perform_comparision(file_docs1,file_docs2):
	lst1 = []
	difference = set(file_docs1).difference(file_docs2)
	difference.discard('\n')
	for l in difference:
		lst1.append(l)
	return lst1;

def find_words(f1,f2):
	file_docs1 = []
	file_docs2 = []
	if(".docx" in f1.name):
		doc = docx.Document(f1)
		for i in doc.paragraphs:
			token1 = word_tokenize(i.text)
			for j in token1 :
				file_docs1.append(j)
	elif(".txt" in f1.name):
		for x in f1:
			temp = x.decode('utf-8')
			temp2 = temp.replace("\n", "").replace("\r", "")
			token1 = word_tokenize(temp2)
			for j in token1:
				file_docs1.append(j)

	if(".docx" in f2.name):
		doc = docx.Document(f2)
		for i in doc.paragraphs:
			token1 = word_tokenize(i.text)
			for j in token1 :
				file_docs2.append(j)
	elif(".txt" in f2.name):
		for x in f2:
			temp = x.decode('utf-8')
			temp2 = temp.replace("\n", "").replace("\r", "")
			token1 = word_tokenize(temp2)
			for j in token1:
				file_docs2.append(j)


	lst1 = []
	difference = set(file_docs1).difference(file_docs2)
	difference.discard('\n')
	for l in difference:
		lst1.append(l)
	print("PARTH:",lst1)
	return lst1;
	

def get_edited_data(f,lst):
	output = ""
	if(".docx" in f.name):
		doc = docx.Document(f)
		for x in doc.paragraphs:
			token1 = sent_tokenize(x.text)
			for i in token1:
				code = ""
				for l in lst:
					if(i == l):
						code = i.replace(l, "<b><mark>" + l + "</mark></b>")	
				if(code != ""):
					output += code
				else:
					output += i
	elif(".txt" in f.name):
		for x in f:
			temp = x.decode('utf-8')
			temp2 = temp.replace("\n", "").replace("\r", "")
			token1 = sent_tokenize(temp2)
			for i in token1:
				code = ""
				for l in lst:
					if(i == l):
						code = i.replace(l, "<b><mark>" + l + "</mark></b>")	
				if(code != ""):
					output += code
				else:
					output += i
	return output
